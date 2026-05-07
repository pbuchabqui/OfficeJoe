"""Service for generating Termo de Diligência DOCX documents."""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.db.models.diligence import Diligence
from app.db.models.diligence_item import DiligenceItem


async def generate_termo_diligencia_docx(
    db: AsyncSession,
    diligence_id: str,
) -> BytesIO:
    """Generate a DOCX file for a Termo de Diligência.

    Args:
        db: Database session
        diligence_id: ID of the diligence

    Returns:
        BytesIO object containing the DOCX file in memory

    Raises:
        ValueError: If diligence not found
    """
    diligence = await db.get(Diligence, diligence_id)
    if not diligence:
        raise ValueError(f"Diligence {diligence_id} not found")

    case = diligence.case
    if not case:
        raise ValueError(f"Case for diligence {diligence_id} not found")

    items = await db.scalars(
        select(DiligenceItem)
        .where(DiligenceItem.diligence_id == diligence_id)
        .order_by(DiligenceItem.created_at.asc())
    )

    current_date = datetime.utcnow().strftime("%d de %B de %Y")
    deadline_formatted = diligence.deadline.strftime("%d de %B de %Y")

    template_env = Environment(
        loader=FileSystemLoader(str(Path(__file__).parent.parent / "templates"))
    )
    template = template_env.get_template("termo_diligencia.jinja2")

    rendered_content = template.render(
        diligence=diligence,
        current_date=current_date,
        deadline_formatted=deadline_formatted,
        case_number=case.case_number,
        case_title=case.title,
        items=items.all(),
    )

    doc = Document()

    for line in rendered_content.split("\n"):
        if line.startswith("TERMO DE DILIGÊNCIA"):
            para = doc.add_paragraph(line)
            para_format = para.paragraph_format
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True

        elif line.startswith("NÚMERO:"):
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True

        elif line.startswith("DATA:"):
            parts = line.split(": ", 1)
            if len(parts) == 2:
                para = doc.add_paragraph()
                run1 = para.add_run(parts[0] + ": ")
                run1.font.size = Pt(11)
                run1.font.bold = True
                run2 = para.add_run(parts[1])
                run2.font.size = Pt(11)

        elif line.startswith("DESTINATÁRIO:"):
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True

        elif line.startswith("PRAZO DE ATENDIMENTO:"):
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True

        elif line.startswith("OBSERVAÇÕES:"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            run = para.add_run(line)
            run.font.size = Pt(11)
            run.font.bold = True

        elif line.startswith("ITENS SOLICITADOS:"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            run = para.add_run(line)
            run.font.size = Pt(11)
            run.font.bold = True

        elif line.startswith("Processo:"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            parts = line.split(": ", 1)
            if len(parts) == 2:
                run1 = para.add_run(parts[0] + ": ")
                run1.font.size = Pt(11)
                run1.font.bold = True
                run2 = para.add_run(parts[1])
                run2.font.size = Pt(11)

        elif line.startswith("Assunto:"):
            parts = line.split(": ", 1)
            if len(parts) == 2:
                para = doc.add_paragraph()
                run1 = para.add_run(parts[0] + ": ")
                run1.font.size = Pt(11)
                run1.font.bold = True
                run2 = para.add_run(parts[1])
                run2.font.size = Pt(11)

        elif line.startswith("Este documento"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(6)
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = para.add_run(line)
            run.font.size = Pt(11)

        elif line.startswith("Local e Data:"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(12)
            run = para.add_run(line)
            run.font.size = Pt(11)

        elif line.startswith("Assinado por:"):
            para = doc.add_paragraph()
            para_format = para.paragraph_format
            para_format.space_before = Pt(24)
            run = para.add_run(line)
            run.font.size = Pt(11)

        elif line.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            para = doc.add_paragraph(line, style="List Number")
            para.paragraph_format.left_indent = Inches(0.5)
            for run in para.runs:
                run.font.size = Pt(11)

        elif line.strip().startswith("Período:"):
            para = doc.add_paragraph(line)
            para.paragraph_format.left_indent = Inches(0.75)
            for run in para.runs:
                run.font.size = Pt(10)

        elif line.strip().startswith("Justificativa Técnica:"):
            para = doc.add_paragraph(line)
            para.paragraph_format.left_indent = Inches(0.75)
            for run in para.runs:
                run.font.size = Pt(10)

        elif line.strip().startswith("Status:"):
            para = doc.add_paragraph(line)
            para.paragraph_format.left_indent = Inches(0.75)
            for run in para.runs:
                run.font.size = Pt(10)

        elif line == "=======================================================================":
            pass

        elif line.strip():
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.size = Pt(11)

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return docx_stream
