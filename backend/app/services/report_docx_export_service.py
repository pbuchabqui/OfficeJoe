"""Simple DOCX export for reports."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from jinja2 import Environment, FileSystemLoader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.case import Case
from app.db.models.report import Report, ReportSection
from app.db.models.report_attachment import ReportAttachment


async def generate_report_docx(db: AsyncSession, report_id: str) -> BytesIO:
    report, case = await _read_report_with_case(db, report_id)
    sections = await _list_sections(db, report_id)
    attachments = await _list_attachments(db, report_id)

    template_env = Environment(
        loader=FileSystemLoader(str(Path(__file__).parent.parent / "templates")),
        autoescape=False,
    )
    template = template_env.get_template("report_docx_basic.jinja2")
    rendered_content = template.render(
        report=report,
        case_number=case.case_number,
        sections=sections,
        attachments=[_attachment_context(item) for item in attachments],
    )

    doc = Document()
    for line in rendered_content.split("\n"):
        text = line.strip()
        if not text:
            continue
        if text == "LAUDO PERICIAL":
            doc.add_heading(text, level=0)
        elif text in {"SEÇÕES", "ANEXOS E APÊNDICES"}:
            doc.add_heading(text, level=1)
        elif _is_ordered_section_heading(text):
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream


async def _read_report_with_case(db: AsyncSession, report_id: str) -> tuple[Report, Case]:
    result = await db.execute(
        select(Report, Case)
        .join(Case, Case.id == Report.case_id)
        .where(Report.id == report_id)
    )
    row = result.one_or_none()
    if not row:
        raise ValueError(f"Report {report_id} not found")
    return row[0], row[1]


async def _list_sections(db: AsyncSession, report_id: str) -> list[ReportSection]:
    result = await db.execute(
        select(ReportSection)
        .where(ReportSection.report_id == report_id)
        .order_by(ReportSection.section_order)
    )
    return result.scalars().all()


async def _list_attachments(db: AsyncSession, report_id: str) -> list[ReportAttachment]:
    result = await db.execute(
        select(ReportAttachment)
        .where(ReportAttachment.report_id == report_id)
        .order_by(ReportAttachment.attachment_type, ReportAttachment.number)
    )
    return result.scalars().all()


def _attachment_context(attachment: ReportAttachment) -> dict[str, str | int | None]:
    label = "Anexo" if attachment.attachment_type == "anexo" else "Apêndice"
    return {
        "label": label,
        "number": attachment.number,
        "title": attachment.title,
        "description": attachment.description,
    }


def _is_ordered_section_heading(text: str) -> bool:
    prefix, separator, suffix = text.partition(". ")
    return bool(separator and prefix.isdigit() and suffix)
