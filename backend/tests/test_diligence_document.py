"""Tests for diligence DOCX document generation."""
from __future__ import annotations

import uuid
from datetime import datetime, timedelta
from io import BytesIO

import pytest
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document

from app.db.models.case import Case, CaseStatus, CaseType
from app.db.models.diligence import Diligence
from app.db.models.diligence_item import DiligenceItem
from app.schemas.diligence import DiligenceCreateRequest, DiligenceItemCreateRequest
from app.services.diligence_service import create_diligence
from app.services.diligence_document_service import generate_termo_diligencia_docx


async def _make_case(db: AsyncSession, user_id: str) -> Case:
    """Helper to create a test case."""
    case = Case(
        id=str(uuid.uuid4()),
        case_number=f"000{uuid.uuid4().hex[:4]}-56.2024.5.02.0000",
        case_type=CaseType.TRABALHISTA.value,
        title="Test Case for Diligence",
        status=CaseStatus.PLANEJAMENTO.value,
        responsible_user_id=user_id,
    )
    db.add(case)
    await db.flush()
    return case


# ── Document Generation Tests ────────────────────────────────────────────


@pytest.mark.asyncio
async def test_generate_termo_diligencia_success(db_session: AsyncSession):
    """Gerar Termo de Diligência DOCX com sucesso."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user1@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-001",
        recipient="Cliente Teste",
        deadline=deadline,
        observations="Documento para testes",
        items=[
            DiligenceItemCreateRequest(
                requested_document="Contrato Principal",
                period="2024-01 a 2024-03",
                technical_justification="Necessário para análise de cláusulas",
            ),
            DiligenceItemCreateRequest(
                requested_document="Recibos de Pagamento",
                period="2024-01 a 2024-03",
                technical_justification="Comprovação de remuneração",
            ),
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    assert docx_stream is not None
    assert isinstance(docx_stream, BytesIO)
    assert docx_stream.getbuffer().nbytes > 0


@pytest.mark.asyncio
async def test_termo_diligencia_contains_diligence_number(db_session: AsyncSession):
    """DOCX gerado contém número da diligência."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user2@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-002",
        recipient="Cliente Teste",
        deadline=deadline,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Documento",
                period="2024-01 a 2024-03",
                technical_justification="Análise",
            )
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    doc = Document(docx_stream)
    full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    assert "2024-DOC-002" in full_text


@pytest.mark.asyncio
async def test_termo_diligencia_contains_recipient(db_session: AsyncSession):
    """DOCX gerado contém destinatário."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user3@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-003",
        recipient="Empresa ACME Ltda",
        deadline=deadline,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Documento",
                period="2024-01 a 2024-03",
                technical_justification="Análise",
            )
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    doc = Document(docx_stream)
    full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    assert "Empresa ACME Ltda" in full_text


@pytest.mark.asyncio
async def test_termo_diligencia_contains_items(db_session: AsyncSession):
    """DOCX gerado contém itens solicitados."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user4@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-004",
        recipient="Cliente Teste",
        deadline=deadline,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Contrato Social",
                period="2024-01 a 2024-03",
                technical_justification="Análise de cláusulas",
            ),
            DiligenceItemCreateRequest(
                requested_document="Balancete Contábil",
                period="2024-01 a 2024-03",
                technical_justification="Comprovação de resultados",
            ),
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    doc = Document(docx_stream)
    full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    assert "Contrato Social" in full_text
    assert "Balancete Contábil" in full_text
    assert "2024-01 a 2024-03" in full_text


@pytest.mark.asyncio
async def test_termo_diligencia_contains_case_info(db_session: AsyncSession):
    """DOCX gerado contém informações do processo."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user5@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-005",
        recipient="Cliente Teste",
        deadline=deadline,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Documento",
                period="2024-01 a 2024-03",
                technical_justification="Análise",
            )
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    doc = Document(docx_stream)
    full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    assert case.case_number in full_text
    assert case.title in full_text


@pytest.mark.asyncio
async def test_termo_diligencia_document_valid(db_session: AsyncSession):
    """Documento DOCX gerado é válido."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user6@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    payload = DiligenceCreateRequest(
        number="2024-DOC-006",
        recipient="Cliente Teste",
        deadline=deadline,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Documento",
                period="2024-01 a 2024-03",
                technical_justification="Análise",
            )
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    try:
        doc = Document(docx_stream)
        assert len(doc.paragraphs) > 0
    except Exception as e:
        pytest.fail(f"Generated DOCX is invalid: {str(e)}")


@pytest.mark.asyncio
async def test_termo_diligencia_not_found(db_session: AsyncSession):
    """Erro ao gerar Termo de Diligência inexistente."""
    with pytest.raises(ValueError, match="Diligence .* not found"):
        await generate_termo_diligencia_docx(db_session, "invalid-id")


@pytest.mark.asyncio
async def test_termo_diligencia_contains_observations(db_session: AsyncSession):
    """DOCX gerado contém observações."""
    from app.db.models.user import User
    from app.core.security import hash_password

    user = User(
        id=str(uuid.uuid4()),
        email="doc-user7@teste.com",
        full_name="Document User",
        hashed_password=hash_password("SenhaSegura123!"),
        role="perito",
        is_active=True,
    )
    db_session.add(user)
    await db_session.flush()

    case = await _make_case(db_session, user.id)

    deadline = datetime.utcnow() + timedelta(days=30)
    observations_text = "Esta é uma solicitação urgente de documentos para fins de perícia"
    payload = DiligenceCreateRequest(
        number="2024-DOC-007",
        recipient="Cliente Teste",
        deadline=deadline,
        observations=observations_text,
        items=[
            DiligenceItemCreateRequest(
                requested_document="Documento",
                period="2024-01 a 2024-03",
                technical_justification="Análise",
            )
        ],
    )

    diligence = await create_diligence(db_session, case.id, payload)
    await db_session.commit()

    docx_stream = await generate_termo_diligencia_docx(db_session, diligence.id)

    doc = Document(docx_stream)
    full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    assert observations_text in full_text
