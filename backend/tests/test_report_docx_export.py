from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.schemas.report import ReportCreateRequest, ReportSectionCreateRequest
from app.schemas.report_attachment import ReportAttachmentCreateRequest
from app.services.report_attachment_service import create_report_attachment
from app.services.report_docx_export_service import generate_report_docx
from app.services.report_service import create_report, create_report_section


@pytest.mark.asyncio
async def test_generate_report_docx_includes_sections_and_attachments(
    db_session: AsyncSession,
    sample_case,
):
    report = await create_report(
        db_session,
        sample_case.id,
        ReportCreateRequest(title="Laudo Trabalhista", report_type="trabalhista"),
    )
    await create_report_section(
        db_session,
        sample_case.id,
        report.id,
        ReportSectionCreateRequest(
            title="Metodologia",
            section_order=2,
            content="Análise documental simples.",
        ),
    )
    await create_report_section(
        db_session,
        sample_case.id,
        report.id,
        ReportSectionCreateRequest(
            title="Objeto",
            section_order=1,
            content="Apuração de verbas.",
        ),
    )
    await create_report_attachment(
        db_session,
        sample_case.id,
        report.id,
        ReportAttachmentCreateRequest(
            attachment_type="anexo",
            title="Holerites analisados",
        ),
    )
    await create_report_attachment(
        db_session,
        sample_case.id,
        report.id,
        ReportAttachmentCreateRequest(
            attachment_type="apendice",
            title="Memória de cálculo",
        ),
    )

    docx_stream = await generate_report_docx(db_session, report.id)

    assert isinstance(docx_stream, BytesIO)
    assert docx_stream.getbuffer().nbytes > 0

    doc = Document(docx_stream)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    text = "\n".join(paragraphs)

    assert "LAUDO PERICIAL" in text
    assert "Título: Laudo Trabalhista" in text
    assert f"Processo: {sample_case.case_number}" in text
    assert paragraphs.index("1. Objeto") < paragraphs.index("2. Metodologia")
    assert "Apuração de verbas." in text
    assert "Anexo 1 - Holerites analisados" in text
    assert "Apêndice 1 - Memória de cálculo" in text


@pytest.mark.asyncio
async def test_report_docx_download_endpoint(
    client: AsyncClient,
    sample_case,
    perito_token: str,
):
    report_response = await client.post(
        f"/api/v1/reports?case_id={sample_case.id}",
        json={"title": "Laudo", "report_type": "trabalhista"},
        headers={"Authorization": f"Bearer {perito_token}"},
    )
    assert report_response.status_code == 201
    report_id = report_response.json()["id"]

    section_response = await client.post(
        f"/api/v1/reports/{report_id}/sections?case_id={sample_case.id}",
        json={"title": "Conclusão", "section_order": 1, "content": "Conclusão preliminar."},
        headers={"Authorization": f"Bearer {perito_token}"},
    )
    assert section_response.status_code == 201

    response = await client.get(
        f"/api/v1/reports/{report_id}/download-docx",
        headers={"Authorization": f"Bearer {perito_token}"},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert f"filename=laudo_{report_id}.docx" in response.headers["content-disposition"]

    doc = Document(BytesIO(response.content))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Conclusão preliminar." in text
